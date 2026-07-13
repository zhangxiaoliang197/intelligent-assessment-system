"""
附件处理模块 — 文档上传解析与持久化缓存
================================================================
职责：
  1. 接收上传的 PDF/Word/TXT 文件，解析为纯文本
  2. 内存缓存 + 磁盘持久化双重存储（data/attachments/）
  3. 服务重启后磁盘数据不丢失，可继续使用
  4. 提供查询接口，供对话流程注入文档上下文

安全措施：
  - 文件大小限制 20MB
  - 仅允许 pdf/doc/docx/txt 格式
  - 文本结果截断至 100KB（防止 token 爆炸）
  - 1 小时 TTL 自动清理（内存 + 磁盘同步清除）
================================================================
"""
import os
import re
import time
import uuid
import json
import hashlib
import logging
import threading
from typing import Optional

logger = logging.getLogger("attachment.handler")

# ─── 配置 ───
MAX_FILE_SIZE = 20 * 1024 * 1024       # 20MB
MAX_TEXT_LENGTH = 100 * 1024            # 100KB 截断
ALLOWED_EXTENSIONS = {"pdf", "doc", "docx", "txt"}
TTL_SECONDS = 3600                      # 1 小时过期

# ─── 持久化存储目录 ───
_STORAGE_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "data", "attachments")
os.makedirs(_STORAGE_DIR, exist_ok=True)

# ─── 内存缓存：{attachment_id: {text, filename, created_at}} ───
_attachment_store: dict = {}
# ─── 内容哈希 → attachment_id 映射（同一文件去重）───
_hash_to_id: dict = {}
_store_lock = threading.Lock()


def _disk_path(attachment_id: str) -> str:
    """获取附件的磁盘持久化路径"""
    return os.path.join(_STORAGE_DIR, f"{attachment_id}.json")


def _save_to_disk(attachment_id: str, data: dict):
    """将附件数据写入磁盘 JSON 文件"""
    try:
        with open(_disk_path(attachment_id), "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False)
    except Exception as e:
        logger.warning(f"Failed to persist attachment {attachment_id} to disk: {e}")


def _load_from_disk(attachment_id: str) -> Optional[dict]:
    """从磁盘加载附件数据"""
    path = _disk_path(attachment_id)
    if not os.path.exists(path):
        return None
    try:
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
        # 过期检查
        if time.time() - data.get("created_at", 0) > TTL_SECONDS:
            _delete_from_disk(attachment_id)
            return None
        # 加载到内存缓存，加速后续访问
        with _store_lock:
            if attachment_id not in _attachment_store:
                _attachment_store[attachment_id] = data
            # 重建哈希映射（服务重启后恢复去重能力）
            fh = data.get("file_hash")
            if fh and fh not in _hash_to_id:
                _hash_to_id[fh] = attachment_id
        return data
    except Exception as e:
        logger.warning(f"Failed to load attachment {attachment_id} from disk: {e}")
        return None


def _delete_from_disk(attachment_id: str):
    """
    从磁盘删除附件的所有相关文件（JSON 元数据 + 原始文件）。

    原始文件以 {attachment_id}_ 为前缀，通过 glob 匹配删除。
    """
    # 删除 JSON 元数据文件
    json_path = _disk_path(attachment_id)
    try:
        if os.path.exists(json_path):
            os.remove(json_path)
    except Exception as e:
        logger.warning(f"Failed to delete attachment JSON {json_path}: {e}")

    # 删除原始文件（匹配 {attachment_id}_ 前缀）
    import glob as glob_mod
    file_pattern = os.path.join(_STORAGE_DIR, f"{attachment_id}_*")
    for orig_file in glob_mod.glob(file_pattern):
        try:
            os.remove(orig_file)
        except Exception:
            pass


# ─── 定时清理（每 5 分钟清理过期条目，内存 + 磁盘同步） ───
def _cleanup_expired():
    """后台线程：每 300 秒清理一次过期附件缓存"""
    while True:
        time.sleep(300)
        with _store_lock:
            now = time.time()
            expired = [
                aid for aid, info in _attachment_store.items()
                if now - info.get("created_at", 0) > TTL_SECONDS
            ]
            for aid in expired:
                info = _attachment_store.get(aid, {})
                fh = info.get("file_hash")
                if fh:
                    _hash_to_id.pop(fh, None)
                _delete_from_disk(aid)
                del _attachment_store[aid]
            if expired:
                logger.info(f"Cleaned {len(expired)} expired attachments, {len(_attachment_store)} remaining")

        # 同时清理磁盘上未被内存引用的过期文件
        try:
            now = time.time()
            for fname in os.listdir(_STORAGE_DIR):
                if not fname.endswith(".json"):
                    continue
                aid = fname[:-5]  # 去掉 .json 后缀
                path = os.path.join(_STORAGE_DIR, fname)
                try:
                    with open(path, "r", encoding="utf-8") as f:
                        data = json.load(f)
                    if now - data.get("created_at", 0) > TTL_SECONDS:
                        os.remove(path)
                except Exception:
                    # 损坏的 JSON 直接删除
                    try:
                        os.remove(path)
                    except Exception:
                        pass
        except Exception as e:
            logger.warning(f"Disk cleanup error: {e}")

_cleanup_thread = threading.Thread(target=_cleanup_expired, daemon=True)
_cleanup_thread.start()


# ═══════════════════════════════════════════════════════════════
# 文件解析函数
# ═══════════════════════════════════════════════════════════════

def _compute_hash(file_path: str) -> str:
    """计算文件内容的 SHA256 哈希，用于去重"""
    sha = hashlib.sha256()
    with open(file_path, 'rb') as f:
        while True:
            chunk = f.read(8192)
            if not chunk:
                break
            sha.update(chunk)
    return sha.hexdigest()


def _read_txt(file_path: str) -> str:
    """读取纯文本文件，自动尝试多种编码"""
    encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
    for enc in encodings:
        try:
            with open(file_path, 'r', encoding=enc) as f:
                return f.read()
        except (UnicodeDecodeError, UnicodeError):
            continue
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    except Exception:
        return ""


def _read_pdf(file_path: str) -> str:
    """使用 PyPDF2 提取 PDF 文本（逐页）"""
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        texts = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(texts)
    except Exception as e:
        logger.warning(f"PDF parsing failed: {e}")
        return f"[PDF 解析失败] {str(e)}"


def _read_docx(file_path: str) -> str:
    """使用 python-docx 提取 Word 文档文本（段落 + 表格）"""
    try:
        from docx import Document
        doc = Document(file_path)
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        # 提取表格内容
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text for cell in row.cells if cell.text.strip()]
                if cells:
                    texts.append(" | ".join(cells))
        return "\n".join(texts)
    except Exception as e:
        logger.warning(f"DOCX parsing failed: {e}")
        return f"[DOCX 解析失败] {str(e)}"


# ═══════════════════════════════════════════════════════════════
# 公共 API
# ═══════════════════════════════════════════════════════════════

def validate_file(filename: str, file_size: int) -> Optional[str]:
    """
    校验文件是否合法。

    Returns:
        错误消息字符串（合法时返回 None）
    """
    ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
    if ext not in ALLOWED_EXTENSIONS:
        return f"不支持的文件格式 .{ext}，仅支持: {', '.join(ALLOWED_EXTENSIONS)}"
    if file_size > MAX_FILE_SIZE:
        return f"文件过大 ({file_size / 1024 / 1024:.1f}MB)，上限 20MB"
    return None


def parse_and_store(file_path: str, filename: str) -> dict:
    """
    解析上传的文件，存入内存缓存 + 磁盘持久化，同时保留原文件副本。

    存储策略：
    - data/attachments/{id}.json  → 解析文本元数据
    - data/attachments/{id}_original{ext} → 原始二进制文件（原样保留）

    Args:
        file_path: 临时文件路径
        filename:  原始文件名

    Returns:
        dict: {attachment_id, filename, text_length, preview}
    """
    # 根据扩展名调用对应的解析函数
    ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''

    if ext == 'pdf':
        text = _read_pdf(file_path)
    elif ext in ('docx', 'doc'):
        text = _read_docx(file_path)
    else:
        text = _read_txt(file_path)

    if not text or text.startswith("[PDF 解析失败]") or text.startswith("[DOCX 解析失败]"):
        raise ValueError(f"文件解析失败，可能为空或格式异常: {text[:100]}")

    # 截断过长文本（与截断后的预览分离）
    original_length = len(text)
    if len(text) > MAX_TEXT_LENGTH:
        text = text[:MAX_TEXT_LENGTH]
        logger.info(f"Text truncated from {original_length} to {MAX_TEXT_LENGTH} chars")

    # 计算文件内容哈希（用于去重，同一文件上传多次只存一份）
    file_hash = _compute_hash(file_path)

    # 去重检查：同一内容已存在时直接返回已有 ID
    with _store_lock:
        existing_id = _hash_to_id.get(file_hash)
    if existing_id:
        # 刷新过期时间
        with _store_lock:
            if existing_id in _attachment_store:
                _attachment_store[existing_id]["created_at"] = time.time()
                # 同步更新磁盘 JSON
                _save_to_disk(existing_id, _attachment_store[existing_id])
        logger.info(f"Duplicate file detected (hash={file_hash[:12]}), reuse attachment {existing_id}")
        # 仍需确定预览和长度
        existing_data = _load_from_disk(existing_id) or {}
        preview = existing_data.get("text", text)[:200].replace('\n', ' ').strip()
        return {
            "attachment_id": existing_id,
            "filename": existing_data.get("filename", filename),
            "text_length": existing_data.get("text_length", len(existing_data.get("text", ""))),
            "preview": preview,
            "duplicate": True,
        }

    # 生成附件 ID
    attachment_id = str(uuid.uuid4())
    preview = text[:200].replace('\n', ' ').strip()

    # 保存原始文件到持久化目录（用 UUID_原文件名 命名，便于识别）
    safe_filename = filename.replace("\\", "_").replace("/", "_").replace(":", "_")
    original_dest = os.path.join(_STORAGE_DIR, f"{attachment_id}_{safe_filename}")
    try:
        with open(file_path, 'rb') as src:
            with open(original_dest, 'wb') as dst:
                dst.write(src.read())
        logger.info(f"Original file saved: {original_dest}")
    except Exception as e:
        logger.warning(f"Failed to save original file: {e}")
        original_dest = ""  # 保存失败时不记录路径

    with _store_lock:
        data = {
            "text": text,
            "filename": filename,
            "text_length": len(text),
            "file_hash": file_hash,
            "original_path": original_dest,
            "created_at": time.time(),
        }
        _attachment_store[attachment_id] = data
        _hash_to_id[file_hash] = attachment_id  # 记录哈希映射用于去重

    # 同步写入磁盘持久化（服务重启后不丢失）
    _save_to_disk(attachment_id, data)

    logger.info(f"Stored attachment {attachment_id}: {filename}, {len(text)} chars (memory + disk + original)")
    return {
        "attachment_id": attachment_id,
        "filename": filename,
        "text_length": len(text),
        "preview": preview,
    }


def get_attachment_text(attachment_id: str) -> Optional[str]:
    """
    根据 attachment_id 获取缓存的解析文本。

    查询顺序：内存 → 磁盘 → None
    服务重启后内存清空，自动从磁盘加载。

    Args:
        attachment_id: 上传时返回的 ID

    Returns:
        解析后的文本字符串，不存在或已过期返回 None
    """
    # 1. 先查内存缓存
    with _store_lock:
        info = _attachment_store.get(attachment_id)

    if info:
        # 惰性过期检查
        if time.time() - info.get("created_at", 0) > TTL_SECONDS:
            with _store_lock:
                fh = info.get("file_hash")
                if fh:
                    _hash_to_id.pop(fh, None)
                _attachment_store.pop(attachment_id, None)
            _delete_from_disk(attachment_id)
            return None
        return info["text"]

    # 2. 内存未命中，从磁盘加载（服务重启后走这条路径）
    disk_data = _load_from_disk(attachment_id)
    return disk_data["text"] if disk_data else None


def get_attachment_original_path(attachment_id: str) -> Optional[str]:
    """
    获取附件的原始文件磁盘路径，供下载端点使用。

    查询顺序：内存缓存 metadata → 磁盘 JSON → None
    同时校验文件是否真实存在。

    Args:
        attachment_id: 上传时返回的 ID

    Returns:
        原始文件的绝对路径，不存在或已过期返回 None
    """
    # 1. 先查内存
    with _store_lock:
        info = _attachment_store.get(attachment_id)

    # 2. 内存未命中，尝试从磁盘加载
    if not info:
        info = _load_from_disk(attachment_id)

    if not info:
        return None

    # 过期检查
    if time.time() - info.get("created_at", 0) > TTL_SECONDS:
        with _store_lock:
            fh = info.get("file_hash")
            if fh:
                _hash_to_id.pop(fh, None)
            _attachment_store.pop(attachment_id, None)
        _delete_from_disk(attachment_id)
        return None

    original_path = info.get("original_path", "")
    if original_path and os.path.exists(original_path):
        return original_path

    # 兜底：original_path 为空或不存，尝试通配匹配
    import glob as glob_mod
    candidates = glob_mod.glob(os.path.join(_STORAGE_DIR, f"{attachment_id}_*"))
    for candidate in candidates:
        if candidate.endswith(".json"):
            continue  # 跳过 JSON 元数据
        if os.path.exists(candidate):
            return candidate

    return None


def get_attachment_filename(attachment_id: str) -> str:
    """
    获取附件的原始文件名。

    Returns:
        原始文件名，不存在返回空字符串
    """
    with _store_lock:
        info = _attachment_store.get(attachment_id)
    if not info:
        info = _load_from_disk(attachment_id)
    return info.get("filename", "") if info else ""
