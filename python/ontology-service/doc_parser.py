"""文档解析模块。

从上传的文档中提取纯文本，支持 PDF / DOCX / TXT / MD / CSV。
复用 knowledge-service 的解析逻辑（PyPDF2 + python-docx），避免跨服务依赖。
"""
import os
import logging

logger = logging.getLogger("ontology-service")


def read_txt(file_path: str) -> str:
    """读取纯文本文件，自动尝试多种编码。"""
    encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
    for enc in encodings:
        try:
            with open(file_path, 'r', encoding=enc) as f:
                return f.read()
        except (UnicodeDecodeError, UnicodeError):
            continue
    # 最后兜底：utf-8 忽略无法解码的字节
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    except Exception as e:
        logger.error(f"读取文本文件失败 {file_path}: {e}")
        return ""


def read_pdf(file_path: str) -> str:
    """使用 PyPDF2 提取 PDF 文本。"""
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        texts = [page.extract_text() for page in reader.pages if page.extract_text()]
        return "\n".join(texts)
    except Exception as e:
        logger.error(f"PDF 解析失败 {file_path}: {e}")
        return f"[PDF解析失败] {str(e)}"


def read_docx(file_path: str) -> str:
    """使用 python-docx 提取 DOCX 文本（含表格）。"""
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # 提取表格内容，用 " | " 拼接单元格
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text for cell in row.cells if cell.text.strip()]
                if cells:
                    tables_text.append(" | ".join(cells))
        return "\n".join(paragraphs + tables_text)
    except Exception as e:
        logger.error(f"DOCX 解析失败 {file_path}: {e}")
        return f"[DOCX解析失败] {str(e)}"


def extract_text(file_path: str, filename: str) -> str:
    """根据文件扩展名分发到对应的解析函数。

    Args:
        file_path: 文件在服务器上的临时路径
        filename: 原始文件名（用于判断扩展名）

    Returns:
        解析后的纯文本
    """
    ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
    if ext in ('txt', 'md', 'csv'):
        return read_txt(file_path)
    elif ext == 'pdf':
        return read_pdf(file_path)
    elif ext in ('docx', 'doc'):
        return read_docx(file_path)
    else:
        # 未知扩展名，按文本尝试
        logger.warning(f"未知文件扩展名 {ext}，按文本尝试解析: {filename}")
        return read_txt(file_path)


def truncate_for_llm(text: str, max_chars: int = 12000) -> str:
    """截断文本以适应 LLM 上下文窗口。

    12000 字符约等于 4000-6000 token，留出空间给 prompt 模板和输出。

    Args:
        text: 原始文本
        max_chars: 最大字符数

    Returns:
        截断后的文本，末尾加省略号
    """
    if len(text) <= max_chars:
        return text
    # 按句子边界截断，避免半句话
    truncated = text[:max_chars]
    for i in range(len(truncated) - 1, max_chars // 2, -1):
        if truncated[i] in '。！？!?\n':
            return truncated[:i + 1] + f"\n\n[... 文档较长，已截断，共 {len(text)} 字符，仅分析前 {i + 1} 字符 ...]"
    return truncated + f"\n\n[... 文档较长，已截断 ...]"
