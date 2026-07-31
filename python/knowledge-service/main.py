from dotenv import load_dotenv, find_dotenv
load_dotenv(find_dotenv(usecwd=True))

from fastapi import FastAPI, UploadFile, File, HTTPException, Form, Query
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List, Optional
from datetime import datetime
import logging
import os
import shutil
import uuid
import json
import re
import tempfile
import traceback
import time
import jieba
from rank_bm25 import BM25Okapi

logging.basicConfig(level=logging.INFO, format='%(asctime)s [%(levelname)s] %(message)s')
logger = logging.getLogger(__name__)

app = FastAPI(
    title="知识库服务 - RAG",
    description="文档上传、解析、分片、语义向量检索（Qdrant + BGE）",
    version="2.0.0"
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ── 路径配置 ──
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOAD_DIR = os.path.join(BASE_DIR, "uploads")
DATA_DIR = os.path.join(BASE_DIR, "data")
MODEL_DIR = os.path.join(BASE_DIR, "models", "bge-small-zh-v1.5")
KNOWLEDGE_FILE = os.path.join(DATA_DIR, "knowledge.json")

os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(DATA_DIR, exist_ok=True)

# ── Qdrant 配置 ──
QDRANT_URL = os.getenv("QDRANT_URL", "http://localhost:6333")
QDRANT_COLLECTION = "knowledge_chunks"

# ── 嵌入模型（延迟加载）──
_embedding_model = None

def get_embedding_model():
    """延迟加载 BGE 嵌入模型（首次调用时加载，之后缓存）"""
    global _embedding_model
    if _embedding_model is None:
        from sentence_transformers import SentenceTransformer
        logger.info(f"正在加载嵌入模型: {MODEL_DIR}")
        _embedding_model = SentenceTransformer(MODEL_DIR)
        logger.info(f"嵌入模型加载完成，向量维度: {_embedding_model.get_sentence_embedding_dimension()}")
    return _embedding_model


# ── 本地 BM25 索引（Qdrant REST 不支持 string-based Prefetch，本地实现双路融合）──
_bm25_index = None
_bm25_chunks = None  # 对应 chunks_db 的快照，用于检测变化

def _tokenize(text: str) -> List[str]:
    """中文分词（jieba 精确模式），返回词列表"""
    return list(jieba.cut(text))

def _ensure_bm25_index():
    """确保 BM25 索引与 chunks_db 同步（懒惰重建）"""
    global _bm25_index, _bm25_chunks
    if _bm25_index is None or _bm25_chunks is not chunks_db:
        if not chunks_db:
            _bm25_index = None
            _bm25_chunks = None
            return
        tokenized = [_tokenize(c.text) for c in chunks_db]
        _bm25_index = BM25Okapi(tokenized)
        _bm25_chunks = chunks_db
        logger.info(f"BM25 索引已构建: {len(tokenized)} 个分片")

def _rrf_fusion(vector_results, bm25_results, k=60):
    """RRF 倒数排序融合：合并两个排序列表，返回按融合分降序的索引列表"""
    scores = {}
    for rank, idx in enumerate(vector_results):
        scores[idx] = scores.get(idx, 0) + 1.0 / (k + rank + 1)
    for rank, idx in enumerate(bm25_results):
        scores[idx] = scores.get(idx, 0) + 1.0 / (k + rank + 1)
    return sorted(scores.keys(), key=lambda x: scores[x], reverse=True)

# ── Qdrant 客户端（延迟连接）──
_qdrant_client = None


def _ensure_payload_indexes(qdrant):
    """创建 payload 索引（幂等：已存在则跳过）"""
    indexes_to_create = [
        ("doc_id", "keyword"),
        ("category", "keyword"),
    ]
    for field_name, field_schema in indexes_to_create:
        try:
            qdrant.create_payload_index(
                collection_name=QDRANT_COLLECTION,
                field_name=field_name,
                field_schema=field_schema,
            )
            logger.info(f"已创建 payload 索引: {field_name} ({field_schema})")
        except Exception as e:
            logger.warning(f"创建 payload 索引 {field_name} 失败: {e}")


def get_qdrant_client():
    """获取 Qdrant 客户端（延迟连接）"""
    global _qdrant_client
    if _qdrant_client is None:
        from qdrant_client import QdrantClient
        from qdrant_client.models import Distance, VectorParams
        logger.info(f"正在连接 Qdrant: {QDRANT_URL}")
        _qdrant_client = QdrantClient(url=QDRANT_URL, timeout=30)
        # 确保集合存在
        if not _qdrant_client.collection_exists(QDRANT_COLLECTION):
            model = get_embedding_model()
            dim = model.get_sentence_embedding_dimension()
            _qdrant_client.create_collection(
                collection_name=QDRANT_COLLECTION,
                vectors_config=VectorParams(size=dim, distance=Distance.COSINE),
            )
            logger.info(f"已创建 Qdrant 集合 '{QDRANT_COLLECTION}'，维度: {dim}")
        # 确保 payload 索引存在（幂等，已存在则跳过）
        _ensure_payload_indexes(_qdrant_client)
    return _qdrant_client


# ── Chunk 数据模型 ──
class Chunk:
    def __init__(self, chunk_id, doc_id, text, title, category, tags, source_file):
        self.chunk_id = chunk_id
        self.doc_id = doc_id
        self.text = text
        self.title = title
        self.category = category
        self.tags = tags
        self.source_file = source_file


# ── 持久化存储（JSON 文件，保留元数据）──
def load_db():
    """从 JSON 文件加载数据，带容错备份恢复"""
    if os.path.exists(KNOWLEDGE_FILE):
        try:
            with open(KNOWLEDGE_FILE, 'r', encoding='utf-8') as f:
                data = json.load(f)
            kdb = data.get('knowledge_db', [])
            cdb = [Chunk(**c) for c in data.get('chunks_db', [])]
            cats = data.get('categories_db', {})
            tags = data.get('tags_db', {})
            logger.info(f"已加载 knowledge.json: {len(kdb)} 文档, {len(cdb)} 分片")
            return kdb, cdb, cats, tags
        except Exception as e:
            logger.error(f"加载 knowledge.json 失败: {e}\n{traceback.format_exc()}")
            backup_file = KNOWLEDGE_FILE + '.bak'
            if os.path.exists(backup_file):
                try:
                    logger.info(f"尝试从备份恢复: {backup_file}")
                    with open(backup_file, 'r', encoding='utf-8') as f:
                        data = json.load(f)
                    kdb = data.get('knowledge_db', [])
                    cdb = [Chunk(**c) for c in data.get('chunks_db', [])]
                    cats = data.get('categories_db', {})
                    tags = data.get('tags_db', {})
                    logger.info(f"从备份恢复: {len(kdb)} 文档, {len(cdb)} 分片")
                    return kdb, cdb, cats, tags
                except Exception as e2:
                    logger.error(f"备份恢复也失败: {e2}")
            logger.warning("以空知识库启动")
    else:
        logger.info(f"未找到 knowledge.json，以空白知识库启动")
    return [], [], {}, {}


def save_db():
    """原子写入：先写临时文件，成功后再替换，失败保留旧文件"""
    data = {
        'knowledge_db': knowledge_db,
        'chunks_db': [vars(c) for c in chunks_db],
        'categories_db': categories_db,
        'tags_db': tags_db
    }
    try:
        json.dumps(data, ensure_ascii=False, default=str)
        dir_path = os.path.dirname(KNOWLEDGE_FILE)
        fd, tmp_path = tempfile.mkstemp(dir=dir_path, suffix='.tmp')
        try:
            with os.fdopen(fd, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2, default=str)
            if os.path.exists(KNOWLEDGE_FILE):
                if os.path.exists(KNOWLEDGE_FILE + '.bak'):
                    os.remove(KNOWLEDGE_FILE + '.bak')
                os.rename(KNOWLEDGE_FILE, KNOWLEDGE_FILE + '.bak')
            os.rename(tmp_path, KNOWLEDGE_FILE)
            logger.info(f"已保存: {len(knowledge_db)} 文档, {len(chunks_db)} 分片")
        except Exception:
            if os.path.exists(tmp_path):
                os.remove(tmp_path)
            raise
    except Exception as e:
        logger.error(f"save_db 失败: {e}\n{traceback.format_exc()}")
        raise


knowledge_db, chunks_db, categories_db, tags_db = load_db()

# ── 嵌入模型信息 ──
EMBEDDING_MODEL_NAME = "BAAI/bge-small-zh-v1.5"
EMBEDDING_MAX_TOKENS = 512  # BGE-small-zh 上下文窗口


# ── 文本分片 ──
DELIMITER = r'[。！？\n;；!?]'


def chunk_text(text: str, chunk_size: int = 400, overlap: int = 100) -> List[str]:
    """将文本按分隔符切分为重叠的分片。

    chunk_size 不宜超过 400（约等于 BGE-small-zh-v1.5 的 512 token 上限）。
    """
    if len(text) <= chunk_size:
        return [text]

    sentences = re.split(f'({DELIMITER})', text)
    merged = []
    buf = ""
    for s in sentences:
        buf += s
        if re.search(DELIMITER, s) and len(buf) >= chunk_size // 2:
            merged.append(buf.strip())
            buf = ""
    if buf.strip():
        merged.append(buf.strip())

    if not merged:
        merged = [text]

    chunks = []
    for seg in merged:
        if len(seg) <= chunk_size:
            chunks.append(seg)
        else:
            start = 0
            while start < len(seg):
                end = start + chunk_size
                chunks.append(seg[start:end])
                start = end - overlap
                if start >= len(seg):
                    break

    return [c for c in chunks if len(c) >= 20]


# ── 文件解析 ──
def read_txt(file_path: str) -> str:
    encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
    for enc in encodings:
        try:
            with open(file_path, 'r', encoding=enc) as f:
                return f.read()
        except (UnicodeDecodeError, UnicodeError):
            continue
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    except Exception:
        return ""


def read_pdf(file_path: str) -> str:
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        texts = [page.extract_text() for page in reader.pages if page.extract_text()]
        return "\n".join(texts)
    except Exception as e:
        return f"[PDF解析失败] {str(e)}"


def read_docx(file_path: str) -> str:
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text for cell in row.cells if cell.text.strip()]
                if cells:
                    tables_text.append(" | ".join(cells))
        return "\n".join(paragraphs + tables_text)
    except Exception as e:
        return f"[DOCX解析失败] {str(e)}"


def extract_text(file_path: str, filename: str) -> str:
    ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
    if ext == 'txt' or ext == 'md' or ext == 'csv':
        return read_txt(file_path)
    elif ext == 'pdf':
        return read_pdf(file_path)
    elif ext in ('docx', 'doc'):
        return read_docx(file_path)
    else:
        return read_txt(file_path)


def get_file_type(filename: str) -> str:
    ext = filename.rsplit('.', 1)[-1].lower()
    return {
        'pdf': 'PDF文档', 'doc': 'Word文档', 'docx': 'Word文档',
        'xls': 'Excel表格', 'xlsx': 'Excel表格',
        'txt': '文本文件', 'md': 'Markdown文档', 'csv': 'CSV数据'
    }.get(ext, '未知类型')


# ── 向量化辅助 ──
def embed_texts(texts: List[str]) -> List[List[float]]:
    """将文本列表转为向量列表"""
    model = get_embedding_model()
    embeddings = model.encode(texts, normalize_embeddings=True, show_progress_bar=False)
    return embeddings.tolist()


def upsert_chunks_to_qdrant(chunks: List[Chunk]):
    """将分片列表写入 Qdrant"""
    if not chunks:
        return
    qdrant = get_qdrant_client()
    texts = [c.text for c in chunks]
    embeddings = embed_texts(texts)
    from qdrant_client.models import PointStruct
    points = [
        PointStruct(
            id=chunk_id_to_int(c.chunk_id),
            vector=emb,
            payload={
                "chunk_id": c.chunk_id,
                "doc_id": c.doc_id,
                "title": c.title,
                "category": c.category,
                "tags": c.tags,
                "source_file": c.source_file,
                "text": c.text,
            }
        )
        for c, emb in zip(chunks, embeddings)
    ]
    qdrant.upsert(collection_name=QDRANT_COLLECTION, points=points)
    logger.info(f"已写入 {len(points)} 个向量到 Qdrant")


def chunk_id_to_int(chunk_id: str) -> int:
    """将 chunk_id（如 kb_abc123_c5）转为整数（用于 Qdrant point id）"""
    return abs(hash(chunk_id)) % (2 ** 63 - 1)


def _chunk_id_to_index(chunk_id: str) -> Optional[int]:
    """根据 chunk_id 查找在 chunks_db 中的索引"""
    for i, c in enumerate(chunks_db):
        if c.chunk_id == chunk_id:
            return i
    return None


def delete_chunks_from_qdrant(doc_id: str):
    """从 Qdrant 删除指定文档的所有分片"""
    from qdrant_client.models import Filter, FieldCondition, MatchValue, FilterSelector
    qdrant = get_qdrant_client()
    qdrant.delete(
        collection_name=QDRANT_COLLECTION,
        points_selector=FilterSelector(
            filter=Filter(
                must=[FieldCondition(key="doc_id", match=MatchValue(value=doc_id))]
            )
        ),
    )
    logger.info(f"已从 Qdrant 删除文档 {doc_id} 的向量")


# ── Pydantic Models ──
class KnowledgeSearchRequest(BaseModel):
    query: str
    top_k: int = 5
    category: Optional[str] = None


# ═══════════════════════ API ═══════════════════════

@app.get("/")
async def root():
    return {"service": "知识库服务", "version": "2.0.0", "status": "running", "engine": "Qdrant + BGE"}


@app.get("/health")
async def health_check():
    return {"status": "healthy", "service": "knowledge-service"}


@app.post("/knowledge/upload")
async def upload_knowledge(
    file: UploadFile = File(...),
    category: str = Form("未分类"),
    tags: str = Form(""),
    chunk_size: int = Form(400),
    chunk_overlap: int = Form(100)
):
    """上传文档：解析 → 分片 → 向量化 → 存入 Qdrant + JSON。

    分片参数：
    - chunk_size: 每个分片最大字符数（默认 500）
    - chunk_overlap: 相邻分片重叠字符数（默认 100）
    """
    try:
        file_id = f"kb_{uuid.uuid4().hex[:8]}"
        safe_name = f"{file_id}_{file.filename}"
        file_path = os.path.join(UPLOAD_DIR, safe_name)

        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        file_size = os.path.getsize(file_path)
        tag_list = [t.strip() for t in tags.split(',') if t.strip()]

        # 解析文本 + 分片（使用可配置参数）
        raw_text = extract_text(file_path, file.filename)
        text_chunks = chunk_text(raw_text, chunk_size=chunk_size, overlap=chunk_overlap)
        if not text_chunks:
            text_chunks = ["[无法解析文档内容]"]

        # 构建 Chunk 对象
        doc_chunks = []
        for i, ctext in enumerate(text_chunks):
            ch = Chunk(
                chunk_id=f"{file_id}_c{i}",
                doc_id=file_id,
                text=ctext,
                title=file.filename.rsplit('.', 1)[0],
                category=category,
                tags=tag_list,
                source_file=file.filename
            )
            chunks_db.append(ch)
            doc_chunks.append(ch)

        # 向量化 + 写入 Qdrant
        upsert_chunks_to_qdrant(doc_chunks)

        # 更新元数据
        item = {
            "id": file_id,
            "title": file.filename.rsplit('.', 1)[0],
            "filename": file.filename,
            "file_path": file_path,
            "file_size": file_size,
            "file_type": get_file_type(file.filename),
            "category": category,
            "tags": tag_list,
            "status": "已完成",
            "upload_time": datetime.now().isoformat(),
            "content_length": len(raw_text),
            "chunk_count": len(text_chunks)
        }
        knowledge_db.append(item)

        if category not in categories_db:
            categories_db[category] = []
        categories_db[category].append(file_id)

        for tag in tag_list:
            if tag not in tags_db:
                tags_db[tag] = []
            tags_db[tag].append(file_id)

        save_db()

        return {
            "success": True,
            "message": f"上传并解析完成 ({len(text_chunks)} 个分片)",
            "data": {
                "id": item["id"],
                "title": item["title"],
                "filename": item["filename"],
                "file_size": item["file_size"],
                "file_type": item["file_type"],
                "category": item["category"],
                "status": item["status"],
                "content_length": item["content_length"],
                "chunk_count": item["chunk_count"]
            }
        }
    except Exception as e:
        logger.error(f"上传失败: {e}\n{traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/knowledge/batch-upload")
async def batch_upload_knowledge(
    files: List[UploadFile] = File(...),
    category: str = Form("未分类")
):
    results = []
    for f in files:
        try:
            result = await upload_knowledge(file=f, category=category, tags="")
            results.append(result)
        except Exception as e:
            results.append({"success": False, "filename": f.filename, "error": str(e)})
    return {"success": True, "total": len(files), "results": results}


@app.post("/knowledge/parse/{knowledge_id}")
async def parse_knowledge(knowledge_id: str):
    """重新解析已上传的文档"""
    for item in knowledge_db:
        if item["id"] == knowledge_id:
            if item["status"] == "已完成":
                return {"success": True, "message": "文档已解析", "data": {"id": knowledge_id, "status": "已完成"}}

            raw_text = extract_text(item["file_path"], item["filename"])
            text_chunks = chunk_text(raw_text)

            # 清理旧分片
            chunks_db[:] = [c for c in chunks_db if c.doc_id != knowledge_id]
            delete_chunks_from_qdrant(knowledge_id)

            doc_chunks = []
            for i, ctext in enumerate(text_chunks):
                ch = Chunk(
                    chunk_id=f"{knowledge_id}_c{i}", doc_id=knowledge_id,
                    text=ctext, title=item["title"], category=item["category"],
                    tags=item["tags"], source_file=item["filename"]
                )
                chunks_db.append(ch)
                doc_chunks.append(ch)

            upsert_chunks_to_qdrant(doc_chunks)

            item["status"] = "已完成"
            item["content_length"] = len(raw_text)
            item["chunk_count"] = len(text_chunks)
            save_db()
            return {"success": True, "message": f"解析完成 ({len(text_chunks)} 个分片)", "data": {"id": knowledge_id, "status": "已完成"}}
    raise HTTPException(status_code=404, detail="文档不存在")


@app.post("/knowledge/search")
async def search_knowledge(request: KnowledgeSearchRequest):
    """混合检索：语义向量（Qdrant） + BM25（本地），RRF 融合排序"""
    if not chunks_db:
        return {"success": True, "query": request.query, "total": 0, "results": []}

    try:
        qdrant = get_qdrant_client()

        # 构建分类过滤条件
        query_filter = None
        if request.category:
            from qdrant_client.models import Filter, FieldCondition, MatchValue
            query_filter = Filter(
                must=[FieldCondition(key="category", match=MatchValue(value=request.category))]
            )

        # ── 路一：语义向量检索 ──
        query_embedding = embed_texts([request.query])[0]
        fetch_limit = min(max(request.top_k * 10, 50), len(chunks_db))
        vector_result = qdrant.query_points(
            collection_name=QDRANT_COLLECTION,
            query=query_embedding,
            limit=fetch_limit,
            query_filter=query_filter,
        )

        # ── 路二：本地 BM25 关键词检索（同步 category 过滤）──
        _ensure_bm25_index()
        tokenized_query = _tokenize(request.query)
        bm25_scores = _bm25_index.get_scores(tokenized_query) if _bm25_index else []
        # 如果指定了分类，BM25 中也排除不匹配的分片
        if request.category:
            bm25_scores = [
                s if chunks_db[i].category == request.category else 0.0
                for i, s in enumerate(bm25_scores)
            ]

        # 取 top_k * 10 作为候选，避免全量传输
        max_candidates = max(request.top_k * 10, 50)
        merge_limit = min(max_candidates, len(chunks_db))

        # 构建向量检索的索引排序列表（截断），同时记录 cosine 相似度
        vector_indices = []
        vector_scores = {}  # idx → cosine 相似度
        seen_v = set()
        for hit in vector_result.points[:merge_limit]:
            payload = hit.payload or {}
            chunk_id = payload.get("chunk_id", "")
            idx = _chunk_id_to_index(chunk_id)
            if idx is not None and idx not in seen_v:
                vector_indices.append(idx)
                vector_scores[idx] = hit.score
                seen_v.add(idx)

        # 构建 BM25 的索引排序列表（按分数降序，截断）
        bm25_ranked = sorted(
            range(len(bm25_scores)), key=lambda i: bm25_scores[i], reverse=True
        )[:merge_limit]

        # ── RRF 融合 ──
        merged_indices = _rrf_fusion(vector_indices, bm25_ranked)

        # ── 后处理：doc_id 去重 + 取 top_k ──
        doc_map = {d["id"]: d for d in knowledge_db}
        results = []
        seen_docs = set()

        for idx in merged_indices:
            chunk = chunks_db[idx]
            if chunk.doc_id in seen_docs:
                continue
            seen_docs.add(chunk.doc_id)

            doc = doc_map.get(chunk.doc_id, {})
            doc_filename = doc.get("filename", "") or chunk.source_file
            display_title = doc_filename.rsplit(".", 1)[0] if doc_filename else chunk.title

            results.append({
                "id": chunk.doc_id,
                "chunk_id": chunk.chunk_id,
                "title": display_title or chunk.title,
                "filename": doc_filename,
                "text": chunk.text,
                "content": chunk.text,
                "score": round(vector_scores.get(idx, 0), 4),
                "category": chunk.category,
                "tags": chunk.tags
            })

            if len(results) >= request.top_k:
                break

        return {"success": True, "query": request.query, "total": len(results), "results": results}

    except Exception as e:
        logger.error(f"混合检索失败: {e}\n{traceback.format_exc()}")
        return {"success": True, "query": request.query, "total": 0, "results": [],
                "message": f"检索异常: {str(e)}"}


@app.get("/knowledge/list")
async def list_knowledge(
    category: Optional[str] = None,
    status: Optional[str] = None,
    page: int = 1,
    page_size: int = 20
):
    filtered = [k for k in knowledge_db]
    if category:
        filtered = [k for k in filtered if k["category"] == category]
    if status:
        filtered = [k for k in filtered if k["status"] == status]
    filtered.sort(key=lambda k: k.get("upload_time", ""), reverse=True)
    total = len(filtered)
    start = (page - 1) * page_size
    items = filtered[start:start + page_size]
    return {"success": True, "total": total, "page": page, "page_size": page_size, "items": items}


@app.get("/knowledge/stats")
async def get_stats():
    total_size = sum(k.get("file_size", 0) for k in knowledge_db)

    # 获取 Qdrant 中的向量数量
    qdrant_count = 0
    try:
        qdrant = get_qdrant_client()
        info = qdrant.get_collection(QDRANT_COLLECTION)
        qdrant_count = info.points_count if info.points_count else 0
    except Exception:
        pass

    return {
        "success": True,
        "data": {
            "total_documents": len(knowledge_db),
            "total_size": total_size,
            "total_size_formatted": f"{total_size / (1024*1024):.2f} MB",
            "total_chunks": len(chunks_db),
            "qdrant_vectors": qdrant_count,
            "categories": len(categories_db),
            "tags": len(tags_db),
            "parsed_documents": len([k for k in knowledge_db if k["status"] == "已完成"]),
            "pending_documents": len([k for k in knowledge_db if k["status"] != "已完成"]),
            "recent_uploads": len([k for k in knowledge_db if (datetime.now() - datetime.fromisoformat(k["upload_time"])).days < 7])
        }
    }


@app.get("/knowledge/categories")
async def get_categories():
    return {"success": True, "total": len(categories_db),
            "categories": [{"name": n, "count": len(ds), "documents": ds} for n, ds in categories_db.items()]}


@app.post("/knowledge/category")
async def create_category(name: str = Form(...)):
    if name in categories_db:
        raise HTTPException(status_code=400, detail="分类已存在")
    categories_db[name] = []
    save_db()
    return {"success": True, "message": f"分类 '{name}' 创建成功"}


@app.delete("/knowledge/category/{category_name}")
async def delete_category(category_name: str):
    if category_name not in categories_db:
        raise HTTPException(status_code=404, detail="分类不存在")
    for item in knowledge_db:
        if item["category"] == category_name:
            item["category"] = "未分类"
            if "未分类" not in categories_db:
                categories_db["未分类"] = []
            categories_db["未分类"].append(item["id"])
    del categories_db[category_name]
    save_db()
    return {"success": True, "message": f"分类 '{category_name}' 删除成功"}


@app.get("/knowledge/tags")
async def get_tags():
    return {"success": True, "total": len(tags_db),
            "tags": [{"name": n, "count": len(ds), "documents": ds} for n, ds in tags_db.items()]}


@app.get("/knowledge/{knowledge_id}")
async def get_knowledge(knowledge_id: str):
    for item in knowledge_db:
        if item["id"] == knowledge_id:
            return {"success": True, "data": item}
    raise HTTPException(status_code=404, detail="文档不存在")


@app.put("/knowledge/{knowledge_id}")
async def update_knowledge(knowledge_id: str, title: str = Form(""), category: str = Form(...), tags: str = Form("")):
    for item in knowledge_db:
        if item["id"] == knowledge_id:
            old_cat = item["category"]
            new_tags = [t.strip() for t in tags.split(',') if t.strip()]
            if title.strip():
                item["title"] = title.strip()
            item["category"] = category
            item["tags"] = new_tags

            if old_cat != category:
                if old_cat in categories_db and knowledge_id in categories_db[old_cat]:
                    categories_db[old_cat].remove(knowledge_id)
                if category not in categories_db:
                    categories_db[category] = []
                categories_db[category].append(knowledge_id)

            # 同步更新 Qdrant 中的 payload
            try:
                from qdrant_client.models import Filter, FieldCondition, MatchValue
                qdrant = get_qdrant_client()
                payload = {"category": category, "tags": new_tags}
                if title.strip():
                    payload["title"] = title.strip()
                qdrant.set_payload(
                    collection_name=QDRANT_COLLECTION,
                    payload=payload,
                    points=Filter(
                        must=[FieldCondition(key="doc_id", match=MatchValue(value=knowledge_id))]
                    ),
                )
            except Exception as e:
                logger.error(f"更新 Qdrant payload 失败: {e}")

            for ch in chunks_db:
                if ch.doc_id == knowledge_id:
                    if title.strip():
                        ch.title = title.strip()
                    ch.category = category
                    ch.tags = new_tags
            save_db()
            return {"success": True, "message": "更新成功"}
    raise HTTPException(status_code=404, detail="文档不存在")


@app.delete("/knowledge/{knowledge_id}")
async def delete_knowledge(knowledge_id: str):
    global chunks_db
    for i, item in enumerate(knowledge_db):
        if item["id"] == knowledge_id:
            # 删除物理文件
            file_path = item.get("file_path", "")
            file_deleted = False
            if file_path and os.path.exists(file_path):
                os.remove(file_path)
                file_deleted = True
                logger.info(f"已删除文件: {file_path}")

            # 清理分类/标签索引
            cat = item.get("category", "")
            if cat in categories_db and knowledge_id in categories_db[cat]:
                categories_db[cat].remove(knowledge_id)
            for t in item.get("tags", []):
                if t in tags_db and knowledge_id in tags_db[t]:
                    tags_db[t].remove(knowledge_id)

            knowledge_db.pop(i)
            chunks_db = [c for c in chunks_db if c.doc_id != knowledge_id]

            # 从 Qdrant 删除向量（失败不阻止本地清理）
            qdrant_deleted = False
            try:
                delete_chunks_from_qdrant(knowledge_id)
                qdrant_deleted = True
            except Exception as e:
                logger.error(f"Qdrant 删除失败，已在本地清理: {e}")

            save_db()

            msg_parts = []
            if file_deleted:
                msg_parts.append("已删除文件")
            if qdrant_deleted:
                msg_parts.append("已清除向量索引")
            return {"success": True, "message": "、".join(msg_parts) if msg_parts else "已移除记录"}
    raise HTTPException(status_code=404, detail="文档不存在")


@app.post("/knowledge/reindex")
async def reindex_all():
    """重建所有文档的向量索引（全量重新向量化并写入 Qdrant）"""
    if not chunks_db:
        return {"success": True, "message": "知识库为空，无需重建"}

    # 清空 Qdrant 集合并重新创建
    qdrant = get_qdrant_client()
    model = get_embedding_model()
    dim = model.get_sentence_embedding_dimension()
    qdrant.delete_collection(QDRANT_COLLECTION)
    qdrant.create_collection(
        collection_name=QDRANT_COLLECTION,
        vectors_config={"size": dim, "distance": "Cosine"},
    )
    _ensure_payload_indexes(qdrant)

    # 分批写入
    batch_size = 100
    total = len(chunks_db)
    for i in range(0, total, batch_size):
        batch = chunks_db[i:i + batch_size]
        upsert_chunks_to_qdrant(batch)
        logger.info(f"重建进度: {min(i + batch_size, total)}/{total}")

    return {"success": True, "message": f"索引重建完成，共 {total} 个向量", "total": total}


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=10252)
